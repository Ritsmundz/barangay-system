from __future__ import annotations

import os
import re
import subprocess
import sys
import textwrap
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.append(str(ROOT))
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

OUT_DIR = ROOT / "artifacts" / "superadmin_manual"
HTML_DIR = OUT_DIR / "html"
IMG_DIR = OUT_DIR / "screenshots"
DOCX_PATH = OUT_DIR / "Superadmin_User_Manual.docx"

SITE_URL = "http://127.0.0.1:8000/admin/"
PROJECT_PATH = str(ROOT)
STATIC_ROOT_URI = (ROOT / "staticfiles").resolve().as_uri()
ADMIN_USERNAME = "manualsuperadmin"
ADMIN_EMAIL = "manualsuperadmin@example.com"
ADMIN_PASSWORD = "ManualPass123!"
VENV_PYTHON = ROOT / "venv" / "Scripts" / "python.exe"


def ensure_dirs() -> None:
    HTML_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)


FONT_SANS = Path(r"C:\Windows\Fonts\arial.ttf")
FONT_SANS_BOLD = Path(r"C:\Windows\Fonts\arialbd.ttf")
FONT_MONO = Path(r"C:\Windows\Fonts\consola.ttf")


def localize_static_paths(html: str) -> str:
    html = re.sub(r'(["\'])/static/', rf"\1{STATIC_ROOT_URI}/", html)
    html = re.sub(r'url\(/static/', f"url({STATIC_ROOT_URI}/", html)
    return html


def export_admin_pages() -> tuple[Path, Path]:
    code = f"""
import os
import sys
from pathlib import Path

root = Path(r"{ROOT}")
sys.path.append(str(root))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings")

import django
django.setup()

from django.contrib.auth import get_user_model
from django.test import Client

html_dir = root / "artifacts" / "superadmin_manual" / "html"
html_dir.mkdir(parents=True, exist_ok=True)

user_model = get_user_model()
user = user_model.objects.filter(username="{ADMIN_USERNAME}").first()
if user is None:
    user = user_model.objects.create_superuser("{ADMIN_USERNAME}", "{ADMIN_EMAIL}", "{ADMIN_PASSWORD}")
else:
    user.email = "{ADMIN_EMAIL}"
    user.is_superuser = True
    user.is_staff = True
    user.set_password("{ADMIN_PASSWORD}")
    user.save()

client = Client()
login_html = client.get("/admin/").content.decode("utf-8")
(html_dir / "admin_login_raw.html").write_text(login_html, encoding="utf-8")

client.force_login(user)
dashboard_html = client.get("/admin/").content.decode("utf-8")
(html_dir / "admin_dashboard_raw.html").write_text(dashboard_html, encoding="utf-8")
"""
    subprocess.run(
        [str(VENV_PYTHON), "-c", code],
        cwd=str(ROOT),
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )

    login_raw = HTML_DIR / "admin_login_raw.html"
    dashboard_raw = HTML_DIR / "admin_dashboard_raw.html"
    login_path = HTML_DIR / "admin_login.html"
    dashboard_path = HTML_DIR / "admin_dashboard.html"
    login_path.write_text(localize_static_paths(login_raw.read_text(encoding="utf-8")), encoding="utf-8")
    dashboard_path.write_text(localize_static_paths(dashboard_raw.read_text(encoding="utf-8")), encoding="utf-8")
    return login_path, dashboard_path


def terminal_html(title: str, subtitle: str, lines: list[str]) -> str:
    rows = "\n".join(
        f'<div class="line{(" output" if line.startswith("OUT>") else "")}">{line}</div>'
        for line in lines
    )
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title}</title>
  <style>
    body {{
      margin: 0;
      font-family: Arial, sans-serif;
      background:
        radial-gradient(circle at top left, rgba(90, 135, 255, 0.22), transparent 28%),
        linear-gradient(180deg, #eef4ff 0%, #dfe9fb 100%);
      color: #0f1c33;
    }}
    .wrap {{
      width: 1280px;
      min-height: 780px;
      margin: 0 auto;
      padding: 48px;
      box-sizing: border-box;
    }}
    .eyebrow {{
      display: inline-block;
      padding: 8px 14px;
      border-radius: 999px;
      background: #e7efff;
      color: #234da8;
      font-size: 14px;
      font-weight: 700;
      letter-spacing: 0.06em;
      text-transform: uppercase;
    }}
    h1 {{
      margin: 18px 0 8px;
      font-size: 34px;
      line-height: 1.05;
    }}
    p {{
      margin: 0 0 28px;
      font-size: 18px;
      line-height: 1.5;
      color: #49607f;
      max-width: 920px;
    }}
    .terminal {{
      border-radius: 22px;
      overflow: hidden;
      box-shadow: 0 24px 60px rgba(16, 31, 62, 0.22);
      border: 1px solid rgba(10, 31, 77, 0.08);
      background: #0e1728;
    }}
    .bar {{
      display: flex;
      align-items: center;
      gap: 8px;
      padding: 16px 18px;
      background: #111f34;
      border-bottom: 1px solid rgba(255,255,255,0.08);
    }}
    .dot {{
      width: 12px;
      height: 12px;
      border-radius: 999px;
      display: inline-block;
    }}
    .red {{ background: #ff5f56; }}
    .yellow {{ background: #ffbd2e; }}
    .green {{ background: #27c93f; }}
    .label {{
      margin-left: 8px;
      color: #c7d3ea;
      font-size: 14px;
      font-weight: 700;
    }}
    .screen {{
      padding: 26px 28px 30px;
      font-family: Consolas, "Courier New", monospace;
      color: #e8eefb;
      font-size: 24px;
      line-height: 1.65;
      white-space: pre-wrap;
    }}
    .line.output {{
      color: #90f7c2;
    }}
  </style>
</head>
<body>
  <div class="wrap">
    <div class="eyebrow">Superadmin Setup</div>
    <h1>{title}</h1>
    <p>{subtitle}</p>
    <div class="terminal">
      <div class="bar">
        <span class="dot red"></span>
        <span class="dot yellow"></span>
        <span class="dot green"></span>
        <span class="label">Windows PowerShell</span>
      </div>
      <div class="screen">{rows}</div>
    </div>
  </div>
</body>
</html>
"""


def create_terminal_pages() -> list[tuple[str, Path]]:
    pages = [
        (
            "Step 1: Open the project in Terminal",
            "Open Windows Terminal or PowerShell, then move into the barangay system project folder.",
            [
                "PS C:\\Users\\Ritsmund> cd C:\\Users\\Ritsmund\\Documents\\barangay_project",
                "OUT> PS C:\\Users\\Ritsmund\\Documents\\barangay_project>",
            ],
            "step_1_terminal.html",
        ),
        (
            "Step 2: Activate the virtual environment",
            "Activate the project virtual environment first so Django uses the correct installed packages.",
            [
                "PS C:\\Users\\Ritsmund\\Documents\\barangay_project> .\\venv\\Scripts\\activate",
                "OUT> (venv) PS C:\\Users\\Ritsmund\\Documents\\barangay_project>",
            ],
            "step_2_activate.html",
        ),
        (
            "Step 3: Create the superadmin account",
            "Run Django's superuser command, then type the username, email, and password for the future superadmin.",
            [
                "(venv) PS C:\\Users\\Ritsmund\\Documents\\barangay_project> python manage.py createsuperuser",
                "Username: superadmin",
                "Email address: superadmin@example.com",
                "Password: ********",
                "Password (again): ********",
                "OUT> Superuser created successfully.",
            ],
            "step_3_createsuperuser.html",
        ),
    ]

    html_files: list[tuple[str, Path]] = []
    for title, subtitle, lines, filename in pages:
        path = HTML_DIR / filename
        path.write_text(terminal_html(title, subtitle, lines), encoding="utf-8")
        html_files.append((filename.replace(".html", ".png"), path))
    return html_files


def font(size: int, *, bold: bool = False, mono: bool = False) -> ImageFont.FreeTypeFont:
    if mono and FONT_MONO.exists():
        return ImageFont.truetype(str(FONT_MONO), size=size)
    path = FONT_SANS_BOLD if bold else FONT_SANS
    if path.exists():
        return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_multiline(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int], max_chars: int, *, fill: str, text_font) -> int:
    x, y = xy
    for block in text.split("\n"):
        wrapped = textwrap.wrap(block, width=max_chars) or [""]
        for line in wrapped:
            draw.text((x, y), line, fill=fill, font=text_font)
            y += text_font.size + 10
    return y


def render_terminal_image(title: str, subtitle: str, lines: list[str], destination: Path) -> None:
    image = Image.new("RGB", (1440, 900), "#e6eefb")
    draw = ImageDraw.Draw(image)

    draw.ellipse((930, -40, 1250, 240), fill="#d7e6ff")
    draw.ellipse((70, 560, 340, 850), fill="#dceafc")

    draw.rounded_rectangle((56, 56, 260, 102), radius=24, fill="#e7efff")
    draw.text((80, 70), "SUPERADMIN SETUP", fill="#244ea8", font=font(18, bold=True))
    draw.text((60, 126), title, fill="#0f1c33", font=font(42, bold=True))
    draw_multiline(draw, subtitle, (60, 186), 70, fill="#4a607f", text_font=font(24))

    outer = (60, 285, 1380, 830)
    draw.rounded_rectangle(outer, radius=28, fill="#0e1728")
    draw.rounded_rectangle((60, 285, 1380, 352), radius=28, fill="#111f34")
    draw.rectangle((60, 320, 1380, 352), fill="#111f34")
    for offset, color in enumerate(["#ff5f56", "#ffbd2e", "#27c93f"]):
        x = 88 + (offset * 24)
        draw.ellipse((x, 310, x + 14, 324), fill=color)
    draw.text((166, 302), "Windows PowerShell", fill="#c7d3ea", font=font(18, bold=True))

    y = 390
    mono_font = font(27, mono=True)
    for line in lines:
        fill = "#90f7c2" if line.startswith("OUT>") else "#e8eefb"
        draw.text((90, y), line, fill=fill, font=mono_font)
        y += 54

    image.save(destination)


def render_admin_login_image(destination: Path) -> None:
    image = Image.new("RGB", (1440, 1100), "#eef3f7")
    draw = ImageDraw.Draw(image)

    draw.rounded_rectangle((40, 30, 1400, 1060), radius=28, fill="#ffffff", outline="#d9e2e8")
    draw.rounded_rectangle((40, 30, 1400, 98), radius=28, fill="#f6f8fa")
    draw.rectangle((40, 70, 1400, 98), fill="#f6f8fa")
    for offset, color in enumerate(["#ff5f56", "#ffbd2e", "#27c93f"]):
        x = 70 + (offset * 24)
        draw.ellipse((x, 56, x + 14, 70), fill=color)
    draw.rounded_rectangle((170, 44, 1190, 82), radius=18, fill="#ffffff", outline="#d0d7de")
    draw.text((192, 52), SITE_URL, fill="#57606a", font=font(20))

    draw.text((130, 165), "Django administration", fill="#417690", font=font(42, bold=True))
    draw.text((130, 228), "Please enter the correct username and password for a staff account.", fill="#4c5f6f", font=font(22))

    draw.text((130, 336), "Username:", fill="#000000", font=font(24, bold=True))
    draw.rounded_rectangle((130, 372, 840, 440), radius=12, outline="#c4c9cc", width=2, fill="#ffffff")
    draw.text((158, 392), ADMIN_USERNAME, fill="#0f1419", font=font(24))

    draw.text((130, 486), "Password:", fill="#000000", font=font(24, bold=True))
    draw.rounded_rectangle((130, 522, 840, 590), radius=12, outline="#c4c9cc", width=2, fill="#ffffff")
    draw.text((158, 542), "***************", fill="#0f1419", font=font(24))

    draw.rounded_rectangle((130, 640, 320, 700), radius=10, fill="#417690")
    draw.text((196, 657), "LOG IN", fill="#ffffff", font=font(22, bold=True))

    note_box = (930, 286, 1298, 700)
    draw.rounded_rectangle(note_box, radius=20, fill="#f6f8fa", outline="#d0d7de")
    draw.text((965, 326), "What to do here", fill="#24292f", font=font(26, bold=True))
    draw_multiline(
        draw,
        "1. Open /admin/.\n2. Enter the superadmin username.\n3. Type the password.\n4. Click LOG IN.",
        (965, 380),
        22,
        fill="#57606a",
        text_font=font(22),
    )

    image.save(destination)


def render_admin_dashboard_image(destination: Path) -> None:
    image = Image.new("RGB", (1440, 1200), "#f3f5f7")
    draw = ImageDraw.Draw(image)

    draw.rounded_rectangle((40, 30, 1400, 1160), radius=28, fill="#ffffff", outline="#d9e2e8")
    draw.rounded_rectangle((40, 30, 1400, 94), radius=28, fill="#417690")
    draw.rectangle((40, 62, 1400, 94), fill="#417690")
    draw.text((80, 48), "Django administration", fill="#ffffff", font=font(34, bold=True))
    draw.text((1118, 52), ADMIN_USERNAME, fill="#e7f1f6", font=font(20))

    draw.text((90, 150), "Site administration", fill="#000000", font=font(34, bold=True))
    draw.text((90, 198), "Use the available models below to manage the barangay system.", fill="#57606a", font=font(22))

    left_cards = [
        ("Authentication and Authorization", ["Users"]),
        ("Residents", ["Audit logs", "Complaints", "Households", "Payments", "Request purposes", "Residents", "Service requests", "Service types", "User profiles"]),
    ]
    right_cards = [
        "Add or edit user accounts",
        "Review resident records",
        "Open household data",
        "Track service requests",
        "Check complaints and payments",
        "View audit logs",
    ]

    y = 270
    for title, items in left_cards:
        height = 130 if len(items) <= 2 else 360
        draw.rounded_rectangle((90, y, 920, y + height), radius=18, fill="#ffffff", outline="#d0d7de")
        draw.rectangle((90, y, 920, y + 48), fill="#79aec8")
        draw.text((112, y + 11), title, fill="#ffffff", font=font(24, bold=True))
        item_y = y + 72
        for item in items:
            draw.text((122, item_y), item, fill="#0f1419", font=font(22))
            draw.text((720, item_y), "Add   Change", fill="#1d70b8", font=font(20))
            item_y += 32
        y += height + 22

    draw.rounded_rectangle((980, 270, 1320, 748), radius=18, fill="#f6f8fa", outline="#d0d7de")
    draw.text((1010, 302), "Typical superadmin tasks", fill="#24292f", font=font(26, bold=True))
    task_y = 360
    for task in right_cards:
        draw.rounded_rectangle((1008, task_y, 1294, task_y + 54), radius=12, fill="#ffffff", outline="#d8dee4")
        draw.text((1030, task_y + 15), task, fill="#57606a", font=font(20))
        task_y += 68

    image.save(destination)


def add_hr(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "DADCE0")
    border.append(bottom)
    p_pr.append(border)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    title = doc.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(26)
    title.font.color.rgb = RGBColor(0, 0, 0)

    for style_name, size in [("Heading 1", 20), ("Heading 2", 16), ("Heading 3", 14)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)


def build_docx(images: dict[str, Path]) -> None:
    doc = Document()
    style_document(doc)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Superadmin User Manual")

    sub = doc.add_paragraph()
    sub.add_run("Barangay E-Governance System\n").bold = True
    sub.add_run("This guide covers how to create and access the Django superadmin account through the terminal and admin panel.")
    add_hr(doc.add_paragraph())

    note = doc.add_paragraph()
    note.add_run("Important: ").bold = True
    note.add_run("The superadmin is created from the terminal using Django's superuser command. It does not use the Captain, Secretary, Treasurer, or Staff role pages.")

    steps = [
        (
            "Step 1: Open the project folder in Terminal",
            "Open Windows Terminal or PowerShell, then go to the project folder where `manage.py` is located.",
            images["step_1_terminal.png"],
        ),
        (
            "Step 2: Activate the virtual environment",
            "Run `\\.venv\\Scripts\\activate` so the project uses the correct Python environment before any Django command.",
            images["step_2_activate.png"],
        ),
        (
            "Step 3: Create the superadmin account",
            "Run `python manage.py createsuperuser`, then provide the desired username, email address, and password. Replace the sample values in this guide with the actual credentials for the future administrator.",
            images["step_3_createsuperuser.png"],
        ),
        (
            "Step 4: Open the Django admin login page",
            "After the account is created, open a browser and visit `http://127.0.0.1:8000/admin/`. This is the correct superadmin entry point for the system.",
            images["admin_login.png"],
        ),
        (
            "Step 5: Sign in and verify access",
            "Log in using the superadmin credentials. If the login is successful, the system opens the Django administration dashboard where the superadmin can manage users and registered records.",
            images["admin_dashboard.png"],
        ),
    ]

    for heading, body, image_path in steps:
        doc.add_paragraph(heading, style="Heading 1")
        paragraph = doc.add_paragraph(body)
        paragraph.paragraph_format.space_after = Pt(8)
        doc.add_picture(str(image_path), width=Inches(6.2))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run(f"Screenshot: {heading}")
        caption_run.italic = True
        caption_run.font.name = "Arial"
        caption_run.font.size = Pt(10)

    doc.add_paragraph("What the superadmin can manage", style="Heading 1")
    for item in [
        "Users and passwords in Django Admin",
        "Residents, households, and service request records",
        "Complaints, payments, audit logs, and user profiles",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Quick reminder", style="Heading 1")
    reminder = doc.add_paragraph()
    reminder.add_run("For future admins: ").bold = True
    reminder.add_run("keep the superadmin password private, use a strong password, and update credentials immediately if the account is shared during turnover.")

    doc.save(DOCX_PATH)


def main() -> None:
    ensure_dirs()
    export_admin_pages()

    image_map: dict[str, Path] = {}
    terminal_specs = [
        (
            "step_1_terminal.png",
            "Step 1: Open the project in Terminal",
            "Open Windows Terminal or PowerShell, then move into the barangay system project folder.",
            [
                "PS C:\\Users\\Ritsmund> cd C:\\Users\\Ritsmund\\Documents\\barangay_project",
                "OUT> PS C:\\Users\\Ritsmund\\Documents\\barangay_project>",
            ],
        ),
        (
            "step_2_activate.png",
            "Step 2: Activate the virtual environment",
            "Activate the project virtual environment first so Django uses the correct installed packages.",
            [
                "PS C:\\Users\\Ritsmund\\Documents\\barangay_project> .\\venv\\Scripts\\activate",
                "OUT> (venv) PS C:\\Users\\Ritsmund\\Documents\\barangay_project>",
            ],
        ),
        (
            "step_3_createsuperuser.png",
            "Step 3: Create the superadmin account",
            "Run Django's superuser command, then type the username, email, and password for the future superadmin.",
            [
                "(venv) PS C:\\Users\\Ritsmund\\Documents\\barangay_project> python manage.py createsuperuser",
                "Username: superadmin",
                "Email address: superadmin@example.com",
                "Password: ********",
                "Password (again): ********",
                "OUT> Superuser created successfully.",
            ],
        ),
    ]

    for image_name, title, subtitle, lines in terminal_specs:
        out_path = IMG_DIR / image_name
        render_terminal_image(title, subtitle, lines, out_path)
        image_map[image_name] = out_path

    login_image = IMG_DIR / "admin_login.png"
    dashboard_image = IMG_DIR / "admin_dashboard.png"
    render_admin_login_image(login_image)
    render_admin_dashboard_image(dashboard_image)
    image_map["admin_login.png"] = login_image
    image_map["admin_dashboard.png"] = dashboard_image

    build_docx(image_map)
    print(f"Created {DOCX_PATH}")


if __name__ == "__main__":
    main()
